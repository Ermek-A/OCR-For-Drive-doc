import re
from PIL import Image
from docx import Document
import tkinter as tk
from tkinter import filedialog, ttk, messagebox, simpledialog
from pathlib import Path
import threading
import json


class DriversLicenseApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Обработчик водительских удостоверений")
        self.root.geometry("1400x800")

        self.files = []
        self.results = []
        self.current_preview = None

        # Пытаемся импортировать доступные OCR библиотеки
        self.ocr_engines = self.detect_ocr_engines()

        self.setup_ui()

    def detect_ocr_engines(self):
        """Определяем какие OCR движки доступны"""
        engines = {}

        # Проверяем EasyOCR
        try:
            import easyocr
            engines['easyocr'] = easyocr.Reader(['ru', 'en'], gpu=False)
            print("✓ EasyOCR загружен")
        except:
            print("✗ EasyOCR не доступен")

        # Проверяем Tesseract
        try:
            import pytesseract
            from PIL import Image
            # Проверяем что Tesseract установлен
            pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            engines['tesseract'] = pytesseract
            print("✓ Tesseract загружен")
        except:
            print("✗ Tesseract не доступен")

        # Проверяем PaddleOCR
        try:
            from paddleocr import PaddleOCR
            engines['paddle'] = PaddleOCR(use_angle_cls=True, lang='ru', use_gpu=False)
            print("✓ PaddleOCR загружен")
        except:
            print("✗ PaddleOCR не доступен")

        return engines

    def setup_ui(self):
        # Заголовок
        header = tk.Frame(self.root, bg="#2563eb", height=100)
        header.pack(fill=tk.X)

        title = tk.Label(
            header,
            text="Универсальный обработчик водительских удостоверений",
            font=("Arial", 20, "bold"),
            bg="#2563eb",
            fg="white"
        )
        title.pack(pady=10)

        subtitle = tk.Label(
            header,
            text=f"Доступно OCR движков: {len(self.ocr_engines)} | С ручной корректировкой",
            font=("Arial", 10),
            bg="#2563eb",
            fg="#dbeafe"
        )
        subtitle.pack()

        # Основной контейнер
        main_container = tk.Frame(self.root, bg="#f0f9ff")
        main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Левая панель - управление
        left_panel = tk.Frame(main_container, bg="#f0f9ff", width=300)
        left_panel.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))

        # Кнопки
        btn_frame = tk.Frame(left_panel, bg="#f0f9ff")
        btn_frame.pack(fill=tk.X, pady=5)

        self.load_btn = tk.Button(
            btn_frame,
            text="📁 Загрузить фото",
            command=self.load_images,
            bg="#3b82f6",
            fg="white",
            font=("Arial", 10, "bold"),
            relief=tk.FLAT,
            padx=15,
            pady=8
        )
        self.load_btn.pack(fill=tk.X, pady=3)

        self.process_btn = tk.Button(
            btn_frame,
            text="⚙️ Авто-обработка",
            command=self.process_images,
            bg="#10b981",
            fg="white",
            font=("Arial", 10, "bold"),
            relief=tk.FLAT,
            padx=15,
            pady=8,
            state=tk.DISABLED
        )
        self.process_btn.pack(fill=tk.X, pady=3)

        self.manual_btn = tk.Button(
            btn_frame,
            text="✏️ Ручной ввод",
            command=self.manual_entry,
            bg="#f59e0b",
            fg="white",
            font=("Arial", 10, "bold"),
            relief=tk.FLAT,
            padx=15,
            pady=8,
            state=tk.DISABLED
        )
        self.manual_btn.pack(fill=tk.X, pady=3)

        self.export_btn = tk.Button(
            btn_frame,
            text="📥 Экспорт в Word",
            command=self.export_to_word,
            bg="#059669",
            fg="white",
            font=("Arial", 10, "bold"),
            relief=tk.FLAT,
            padx=15,
            pady=8,
            state=tk.DISABLED
        )
        self.export_btn.pack(fill=tk.X, pady=3)

        self.clear_btn = tk.Button(
            btn_frame,
            text="🗑️ Очистить",
            command=self.clear_all,
            bg="#ef4444",
            fg="white",
            font=("Arial", 10, "bold"),
            relief=tk.FLAT,
            padx=15,
            pady=8
        )
        self.clear_btn.pack(fill=tk.X, pady=3)

        # Счетчик
        self.counter_label = tk.Label(
            left_panel,
            text="Файлов: 0",
            font=("Arial", 11, "bold"),
            bg="#f0f9ff"
        )
        self.counter_label.pack(pady=10)

        # Прогресс
        self.progress = ttk.Progressbar(left_panel, mode='determinate')
        self.progress.pack(fill=tk.X, pady=5)

        # Информация о OCR
        info_frame = tk.LabelFrame(left_panel, text="Доступные OCR", bg="#f0f9ff")
        info_frame.pack(fill=tk.X, pady=10)

        if self.ocr_engines:
            for engine in self.ocr_engines.keys():
                tk.Label(
                    info_frame,
                    text=f"✓ {engine.upper()}",
                    bg="#f0f9ff",
                    fg="#059669",
                    font=("Arial", 9)
                ).pack(anchor=tk.W, padx=5, pady=2)
        else:
            tk.Label(
                info_frame,
                text="⚠️ OCR не найден",
                bg="#f0f9ff",
                fg="#ef4444",
                font=("Arial", 9, "bold")
            ).pack(padx=5, pady=5)
            tk.Label(
                info_frame,
                text="Используйте ручной ввод",
                bg="#f0f9ff",
                font=("Arial", 8)
            ).pack(padx=5)

        # Правая панель - таблица
        right_panel = tk.Frame(main_container, bg="white")
        right_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Таблица
        table_frame = tk.Frame(right_panel, bg="white", relief=tk.SOLID, borderwidth=1)
        table_frame.pack(fill=tk.BOTH, expand=True)

        scrollbar_y = tk.Scrollbar(table_frame)
        scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)

        scrollbar_x = tk.Scrollbar(table_frame, orient=tk.HORIZONTAL)
        scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)

        self.tree = ttk.Treeview(
            table_frame,
            columns=("num", "fio", "birth", "license", "issue", "organ", "category"),
            show="headings",
            yscrollcommand=scrollbar_y.set,
            xscrollcommand=scrollbar_x.set
        )

        scrollbar_y.config(command=self.tree.yview)
        scrollbar_x.config(command=self.tree.xview)

        columns = {
            "num": ("№", 50),
            "fio": ("ФИО", 200),
            "birth": ("Туулган күнү", 110),
            "license": ("АК номери", 120),
            "issue": ("АК берилген дата", 120),
            "organ": ("АК берген орган", 180),
            "category": ("Категориясы", 100)
        }

        for col, (heading, width) in columns.items():
            self.tree.heading(col, text=heading)
            self.tree.column(col, width=width, anchor=tk.CENTER if col == "num" else tk.W)

        self.tree.pack(fill=tk.BOTH, expand=True)

        # Двойной клик для редактирования
        self.tree.bind('<Double-1>', self.edit_row)

        # Статус-бар
        self.status_label = tk.Label(
            self.root,
            text="Готов к работе | Совет: используйте ручной ввод для точности",
            font=("Arial", 9),
            bg="#e5e7eb",
            anchor=tk.W,
            padx=10
        )
        self.status_label.pack(fill=tk.X, side=tk.BOTTOM)

    def load_images(self):
        files = filedialog.askopenfilenames(
            title="Выберите изображения",
            filetypes=[
                ("Изображения", "*.jpg *.jpeg *.png *.bmp"),
                ("Все файлы", "*.*")
            ]
        )

        if files:
            self.files.extend(files)
            self.counter_label.config(text=f"Файлов: {len(self.files)}")
            self.process_btn.config(state=tk.NORMAL)
            self.manual_btn.config(state=tk.NORMAL)
            self.status_label.config(text=f"Загружено файлов: {len(files)}")

    def clear_all(self):
        self.files = []
        self.results = []
        self.tree.delete(*self.tree.get_children())
        self.counter_label.config(text="Файлов: 0")
        self.process_btn.config(state=tk.DISABLED)
        self.manual_btn.config(state=tk.DISABLED)
        self.export_btn.config(state=tk.DISABLED)
        self.progress['value'] = 0
        self.status_label.config(text="Готов к работе")

    def process_images(self):
        if not self.files:
            messagebox.showwarning("Предупреждение", "Не загружены файлы")
            return

        if not self.ocr_engines:
            messagebox.showinfo(
                "OCR не доступен",
                "OCR движки не найдены.\n\nИспользуйте 'Ручной ввод' для заполнения данных."
            )
            return

        thread = threading.Thread(target=self._process_thread)
        thread.start()

    def _process_thread(self):
        self.root.after(0, lambda: self.process_btn.config(state=tk.DISABLED))
        self.root.after(0, lambda: self.load_btn.config(state=tk.DISABLED))

        total = len(self.files)

        for i, image_path in enumerate(self.files):
            try:
                progress = (i + 1) / total * 100
                self.root.after(0, lambda p=progress: self.progress.config(value=p))
                self.root.after(0, lambda n=Path(image_path).name:
                self.status_label.config(text=f"Обработка: {n}"))

                data = self.extract_with_available_ocr(image_path)
                self.results.append(data)

                self.root.after(0, lambda d=data, idx=i + 1: self.add_to_table(d, idx))

            except Exception as e:
                error_data = {
                    "ФИО": "ОШИБКА OCR",
                    "Туулган күнү": "",
                    "АК номери": "",
                    "АК берилген дата": "",
                    "АК берген орган": "",
                    "Категориясы": ""
                }
                self.results.append(error_data)
                self.root.after(0, lambda d=error_data, idx=i + 1: self.add_to_table(d, idx))

        self.root.after(0, lambda: self.process_btn.config(state=tk.NORMAL))
        self.root.after(0, lambda: self.load_btn.config(state=tk.NORMAL))
        self.root.after(0, lambda: self.export_btn.config(state=tk.NORMAL))
        self.root.after(0, lambda: messagebox.showinfo(
            "Завершено",
            f"Обработано: {total}\n\nПроверьте результаты и исправьте при необходимости\n(двойной клик на строке)"
        ))

    def extract_with_available_ocr(self, image_path):
        """Используем доступный OCR движок"""
        text = ""

        if 'easyocr' in self.ocr_engines:
            results = self.ocr_engines['easyocr'].readtext(image_path)
            text = "\n".join([t for (bbox, t, prob) in results])

        elif 'tesseract' in self.ocr_engines:
            img = Image.open(image_path)
            text = self.ocr_engines['tesseract'].image_to_string(img, lang='rus+eng')

        elif 'paddle' in self.ocr_engines:
            result = self.ocr_engines['paddle'].ocr(image_path, cls=True)
            text = "\n".join([line[1][0] for line in result[0]])

        return self.parse_text_to_data(text)

    def parse_text_to_data(self, text):
        """Универсальный парсер текста"""
        data = {
            "ФИО": "",
            "Туулган күнү": "",
            "АК номери": "",
            "АК берилген дата": "",
            "АК берген орган": "",
            "Категориясы": ""
        }

        # ФИО - ищем слова на кириллице
        fio_words = re.findall(r'[А-ЯЁ][а-яё]{2,}', text)
        if len(fio_words) >= 3:
            data["ФИО"] = " ".join(fio_words[:3])

        # Даты
        dates = re.findall(r'\d{2}[.\s]\d{2}[.\s]\d{4}', text)
        if len(dates) >= 1:
            data["Туулган күнү"] = dates[0].replace(' ', '.')
        if len(dates) >= 2:
            data["АК берилген дата"] = dates[1].replace(' ', '.')

        # Номер - 8-9 цифр
        numbers = re.findall(r'\b\d{8,9}\b', text)
        if numbers:
            data["АК номери"] = numbers[0]

        # Орган выдачи - ищем код и преобразуем
        organ_codes = {
            # Короткие коды
            "AA": "г. Баткен",
            "AC": "г. Сулюкта",
            "AK": "г. Кызыл-Кыя",
            "BG": "г. Бишкек",
            "CK": "Склад",
            "DA": "Аксыйскому, Ала-Букинскому, Чаткальскому р-нам",
            "DG": "г. Джалал-Абад",
            "DK": "г. Каракуль",
            "DT": "Ноокенскому р-ну",
            "DU": "пгт. Токтогул",
            "GO": "г. Ош",
            "IB": "г. Балыкчы",
            "IK": "г. Каракол",
            "NA": "Ак-Талинскому, Тогуз-Тороузскому р-нам",
            "NG": "г. Нарын",
            "NH": "Джумгальского р-на",
            "NK": "Кочкорского р-на",
            "OG": "Ошское региональное РЭО",
            "OO": "г. Ош",
            "OU": "г. Узген",
            "SA": "Аламудунского р-на",
            "SC": "Сокулукскому, Московскому р-нам",
            "SH": "Жайылскому и Панфиловскому р-нам",
            "SK": "Иссык-Атинскому р-ну",
            "ST": "Чуй-Кеминскому р-нам",
            "TG": "г. Талас",
            "TS": "г. Талас",
            "UG": "Центральный аппарат",

            # MKK / МКК коды
            "MKK 411011": "Бишкекский городской РЭО",
            "MKK 411021": "Восточный отдел",
            "MKK 412011": "Ошский городской РЭО",
            "MKK 413011": "Кызыл-Кийский РЭП Баткенского регионального РЭО",
            "MKK 413021": "Баткенский региональный РЭО",
            "MKK 413031": "Сулюктинский РЭГ Баткенского регионального РЭО",
            "MKK 414011": "Джалал-Абадский региональный РЭО",
            "MKK 414021": "Аксыйский РЭП Джалал-Абадского РЭО",
            "MKK 414031": "Кара-Кульский РЭП Джалал-Абадского РЭО",
            "MKK 414041": "Ноокенский РЭП Джалал-Абадского РЭО",
            "MKK 414051": "Токтогульский РЭП Джалал-Абадского РЭО",
            "MKK 415011": "Нарынский региональный РЭО",
            "MKK 415021": "Ак-Талинский РЭП Нарынского регионального РЭО",
            "MKK 415031": "Джумгальский РЭП Нарынского регионального РЭО",
            "MKK 415041": "Кочкорский РЭП Нарынского регионального РЭО",
            "MKK 416011": "Ошский региональный РЭО",
            "MKK 416021": "Алайский РЭП Ошского регионального РЭО",
            "MKK 416031": "Узгенский РЭП Ошского регионального РЭО",
            "MKK 417011": "Таласский региональный РЭО",
            "MKK 417021": "Кара-Бууринский РЭП Таласского регионального РЭО",
            "MKK 418011": "Аламудунский РЭО",
            "MKK 418021": "Жайыл-Панфиловский РЭО",
            "MKK 418031": "Сокулук-Московский РЭО",
            "MKK 418041": "Чуй-Кеминский РЭО",
            "MKK 418051": "Ысык-Атинский РЭО",
            "MKK 418061": "Межрегиональный отдел по первичной регистрации транспортных средств",
            "MKK 419011": "Иссык-Кульский региональный РЭО",
            "MKK 419021": "Балыкчинский РЭП Иссык-Кульского регионального РЭО",
        }

        # Ищем коды в тексте
        for code, organ_name in organ_codes.items():
            if code in text.upper():
                data["АК берген орган"] = organ_name
                break

        # Категория - улучшенная логика с приоритетом на поле 9
        lines = text.split('\n')

        found_categories = []
        category_text = ""

        # Ищем строку с полем 9 (категории)
        for i, line in enumerate(lines):
            # Ищем строку, которая начинается с "9" или содержит "9."
            if re.match(r'^\s*9[\.\)\s:]', line) or re.search(r'\s9[\.\)\s:]', line):
                # Нормализуем кириллицу в латиницу
                normalized_line = (
                    line.upper()
                    .replace('А', 'A')
                    .replace('В', 'B')
                    .replace('С', 'C')
                    .replace('Д', 'D')
                )

                # Извлекаем часть после "9"
                match = re.search(r'9[\.\)\s:]\s*(.+)', normalized_line)
                if match:
                    category_text = match.group(1).strip()
                    break

        # Если нашли строку с категориями
        if category_text:
            # Ищем паттерны типа "B-B1", "B-C", "A-B-C-D" и т.д.
            # Сначала пробуем найти категории через дефис или пробел
            category_matches = re.findall(r'([ABCD]1?(?:E)?)', category_text)

            # Список допустимых категорий
            valid_categories = ['A1', 'A', 'B1', 'B', 'C1E', 'C1', 'CE', 'C', 'D1E', 'D1', 'DE', 'D']

            # Фильтруем только валидные категории
            for match in category_matches:
                if match in valid_categories and match not in found_categories:
                    found_categories.append(match)

        # Если не нашли через поле 9, ищем в нижней части (как запасной вариант)
        if not found_categories:
            # Ищем в последних строках документа
            last_lines = '\n'.join(lines[-8:])
            normalized_text = (
                last_lines.upper()
                .replace('А', 'A')
                .replace('В', 'B')
                .replace('С', 'C')
                .replace('Д', 'D')
            )

            # Ищем отдельно стоящие категории
            valid_categories = ['A1', 'A', 'B1', 'B', 'C1E', 'C1', 'CE', 'C', 'D1E', 'D1', 'DE', 'D']

            for cat in valid_categories:
                # Ищем категорию как отдельное слово с учетом возможного дефиса
                if re.search(r'(?:^|[\s\-])' + re.escape(cat) + r'(?:[\s\-]|$)', normalized_text):
                    if cat not in found_categories:
                        found_categories.append(cat)

        # Сортируем категории в правильном порядке
        if found_categories:
            category_order = ['A1', 'A', 'B1', 'B', 'C1', 'C1E', 'C', 'CE', 'D1', 'D1E', 'D', 'DE']
            sorted_categories = sorted(found_categories,
                                       key=lambda x: category_order.index(x) if x in category_order else 999)

            # Форматируем вывод
            if len(sorted_categories) == 1:
                data["Категориясы"] = sorted_categories[0]
            else:
                data["Категориясы"] = "-".join(sorted_categories)

        return data
    def manual_entry(self):
        """Ручной ввод данных с просмотром изображения"""
        if not self.files:
            messagebox.showwarning("Предупреждение", "Сначала загрузите файлы")
            return

        # Окно для ручного ввода
        manual_window = tk.Toplevel(self.root)
        manual_window.title("Ручной ввод данных")
        manual_window.geometry("900x600")

        current_file_idx = [0]  # Используем список для изменяемости

        # Левая часть - изображение
        left_frame = tk.Frame(manual_window)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        img_label = tk.Label(left_frame, text="Изображение загружается...")
        img_label.pack(fill=tk.BOTH, expand=True)

        # Правая часть - форма
        right_frame = tk.Frame(manual_window)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=10, pady=10)

        tk.Label(right_frame, text="Введите данные:", font=("Arial", 12, "bold")).pack(pady=5)

        fields = {}
        field_names = [
            ("ФИО", "Досалиева Салима Тузановна"),
            ("Туулган күнү", "26.05.1970"),
            ("АК номери", "000583746"),
            ("АК берилген дата", "11.08.2015"),
            ("АК берген орган", "Бишкекское ГУ"),
            ("Категориясы", "B")
        ]

        for label, placeholder in field_names:
            tk.Label(right_frame, text=label, font=("Arial", 10)).pack(anchor=tk.W, pady=(10, 0))
            entry = tk.Entry(right_frame, font=("Arial", 10), width=30)
            entry.pack(fill=tk.X, pady=(0, 5))
            entry.insert(0, placeholder)
            entry.config(fg='gray')

            # Placeholder эффект
            def on_focus_in(e, ent=entry, ph=placeholder):
                if ent.get() == ph:
                    ent.delete(0, tk.END)
                    ent.config(fg='black')

            def on_focus_out(e, ent=entry, ph=placeholder):
                if not ent.get():
                    ent.insert(0, ph)
                    ent.config(fg='gray')

            entry.bind('<FocusIn>', on_focus_in)
            entry.bind('<FocusOut>', on_focus_out)

            fields[label] = entry

        file_label = tk.Label(right_frame, text="", font=("Arial", 9), wraplength=250)
        file_label.pack(pady=10)

        def load_current_image():
            if current_file_idx[0] < len(self.files):
                img_path = self.files[current_file_idx[0]]
                file_label.config(text=f"Файл {current_file_idx[0] + 1}/{len(self.files)}:\n{Path(img_path).name}")

                try:
                    img = Image.open(img_path)
                    img.thumbnail((500, 500))

                    from PIL import ImageTk
                    photo = ImageTk.PhotoImage(img)
                    img_label.config(image=photo, text="")
                    img_label.image = photo
                except:
                    img_label.config(text="Ошибка загрузки изображения")

        def save_and_next():
            data = {}
            for label, entry in fields.items():
                value = entry.get()
                # Игнорируем placeholder
                if value and value != field_names[[f[0] for f in field_names].index(label)][1]:
                    data[label] = value
                else:
                    data[label] = ""

            self.results.append(data)
            self.add_to_table(data, len(self.results))

            current_file_idx[0] += 1
            if current_file_idx[0] < len(self.files):
                load_current_image()
                # Очищаем поля
                for entry in fields.values():
                    entry.delete(0, tk.END)
            else:
                messagebox.showinfo("Завершено", "Все файлы обработаны!")
                self.export_btn.config(state=tk.NORMAL)
                manual_window.destroy()

        btn_frame = tk.Frame(right_frame)
        btn_frame.pack(pady=20)

        tk.Button(
            btn_frame,
            text="💾 Сохранить и Далее",
            command=save_and_next,
            bg="#10b981",
            fg="white",
            font=("Arial", 10, "bold"),
            padx=20,
            pady=10
        ).pack()

        load_current_image()

    def edit_row(self, event):
        """Редактирование строки по двойному клику"""
        item = self.tree.selection()
        if not item:
            return

        item = item[0]
        values = self.tree.item(item)['values']

        # Окно редактирования
        edit_window = tk.Toplevel(self.root)
        edit_window.title("Редактирование записи")
        edit_window.geometry("400x400")

        fields = {}
        field_names = ["ФИО", "Туулган күнү", "АК номери", "АК берилген дата", "АК берген орган", "Категориясы"]

        for i, label in enumerate(field_names):
            tk.Label(edit_window, text=label, font=("Arial", 10)).pack(anchor=tk.W, padx=20, pady=(10, 0))
            entry = tk.Entry(edit_window, font=("Arial", 10), width=40)
            entry.pack(padx=20, pady=(0, 5))
            entry.insert(0, values[i + 1])  # +1 т.к. первое значение - номер
            fields[label] = entry

        def save_changes():
            new_values = [values[0]] + [fields[f].get() for f in field_names]
            self.tree.item(item, values=new_values)

            # Обновляем в results
            idx = int(values[0].replace('.', '')) - 1
            if 0 <= idx < len(self.results):
                for i, field in enumerate(field_names):
                    self.results[idx][field] = fields[field].get()

            edit_window.destroy()

        tk.Button(
            edit_window,
            text="💾 Сохранить",
            command=save_changes,
            bg="#10b981",
            fg="white",
            font=("Arial", 10, "bold"),
            padx=30,
            pady=10
        ).pack(pady=20)

    def add_to_table(self, data, index):
        """Добавление данных в таблицу"""
        self.tree.insert("", tk.END, values=(
            f"{index}.",
            data.get("ФИО", ""),
            data.get("Туулган күнү", ""),
            data.get("АК номери", ""),
            data.get("АК берилген дата", ""),
            data.get("АК берген орган", ""),
            data.get("Категориясы", "")
        ))

    def export_to_word(self):
        if not self.results:
            messagebox.showwarning("Предупреждение", "Нет данных для экспорта")
            return

        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word документ", "*.docx")]
        )

        if not filename:
            return

        try:
            doc = Document()
            doc.add_heading("Список водительских удостоверений", level=1)

            table = doc.add_table(rows=1, cols=8)
            table.style = "Table Grid"

            headers = [
                "Катер №", "ФАА", "Туулган күнү",
                "АК номери", "АК берилген дата",
                "АК берген орган", "Категориясы",
                "Примечание"
            ]

            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h

            for idx, data in enumerate(self.results, 1):
                row = table.add_row().cells
                row[0].text = f"{idx}."
                row[1].text = data.get("ФИО", "")
                row[2].text = data.get("Туулган күнү", "")
                row[3].text = data.get("АК номери", "")
                row[4].text = data.get("АК берилген дата", "")
                row[5].text = data.get("АК берген орган", "")
                row[6].text = data.get("Категориясы", "")
                row[7].text = ""

            doc.save(filename)
            messagebox.showinfo("Успех", f"Документ сохранен:\n{filename}")
            self.status_label.config(text=f"Экспортировано в: {Path(filename).name}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте:\n{str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = DriversLicenseApp(root)
    root.mainloop()